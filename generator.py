"""
Resume & Cover Letter Generator — Phase 2: AI-Powered Document Generation

Takes the top-ranked jobs from the matcher and generates tailored
resumes and cover letters as DOCX files using Anthropic Claude.

All AI settings, prompts, and personal info are loaded from ai_config.yaml.

Usage:
    python generator.py                             # top 5 from latest best_matches
    python generator.py --top 10                    # generate for top 10 jobs
    python generator.py --matches best_matches.csv  # specific matches file
    python generator.py --jobs jobs.csv             # specific jobs file
    python generator.py --config my_config.yaml     # use a different config

Requires:
    ANTHROPIC_API_KEY environment variable (or .env file)
"""

import csv
import os
import re
import glob
import argparse
from datetime import datetime

import yaml
from anthropic import Anthropic
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ──────────────────────────────────────────────────────────────────────
# CONFIG LOADING
# ──────────────────────────────────────────────────────────────────────

DEFAULT_CONFIG_PATH = "ai_config.yaml"


def load_config(config_path: str) -> dict:
    """Load AI configuration from YAML file."""
    if not os.path.exists(config_path):
        print(f"❌ Config file not found: {config_path}")
        print(f"   Create one from the template or use --config to specify a path")
        raise FileNotFoundError(config_path)

    with open(config_path, "r", encoding="utf-8") as f:
        config = yaml.safe_load(f)

    return config


def load_base_resume(resume_path: str = "base_resume.md") -> str:
    """Load the base resume from a text file."""
    if not os.path.exists(resume_path):
        print(f"❌ Base resume file not found: {resume_path}")
        print(f"   Create a base_resume.md file with your resume in Markdown")
        raise FileNotFoundError(resume_path)

    with open(resume_path, "r", encoding="utf-8") as f:
        return f.read().strip()


# ──────────────────────────────────────────────────────────────────────
# FILE DISCOVERY
# ──────────────────────────────────────────────────────────────────────

def find_latest_file(pattern: str) -> str | None:
    """Find the most recent file matching a glob pattern by modification time."""
    files = glob.glob(pattern)
    if files:
        files.sort(key=os.path.getmtime, reverse=True)
        return files[0]
    return None


def find_latest_matches_csv() -> str | None:
    """Find the most recent best_matches CSV."""
    result = find_latest_file("best_matches_*.csv")
    if result:
        return result
    if os.path.exists("best_matches.csv"):
        return "best_matches.csv"
    return None


def find_latest_jobs_csv() -> str | None:
    """Find the most recent jobs CSV (for full descriptions)."""
    result = find_latest_file("jobs_*.csv")
    if result:
        return result
    if os.path.exists("jobs.csv"):
        return "jobs.csv"
    return None


# ──────────────────────────────────────────────────────────────────────
# DATA LOADING
# ──────────────────────────────────────────────────────────────────────

def load_csv(file_path: str) -> list[dict[str, str]]:
    """Load rows from a CSV file."""
    rows: list[dict[str, str]] = []
    with open(file_path, mode="r", encoding="utf-8") as f:
        reader = csv.DictReader(f)
        for row in reader:
            rows.append(row)
    return rows


def build_job_description_map(jobs_csv: str) -> dict[str, str]:
    """Build a map of job_link -> job_description from the jobs CSV."""
    jobs = load_csv(jobs_csv)
    return {
        row.get("job_link", "").strip(): row.get("job_description", "")
        for row in jobs
        if row.get("job_link", "").strip()
    }


# ──────────────────────────────────────────────────────────────────────
# AI GENERATION
# ──────────────────────────────────────────────────────────────────────

def generate_tailored_resume(
    client: Anthropic,
    config: dict,
    base_resume: str,
    job_title: str,
    company_name: str,
    job_description: str,
) -> str:
    """Use Claude to generate a tailored resume for a specific job."""
    global_rules = config.get("global_instructions", "")
    instructions = config.get("resume_instructions", "Tailor the resume for the job.")
    output_format = config.get("resume_output_format", "Return only the tailored resume text.")
    model = config.get("model", "claude-sonnet-4-20250514")
    max_tokens = config.get("resume_max_tokens", 2000)
    temperature = config.get("temperature", 1.0)

    prompt = f"""You are a professional resume writer. Your task is to tailor the following base resume for a specific job application.

GLOBAL RULES (always follow these):
{global_rules}

BASE RESUME:
{base_resume}

TARGET JOB:
Title: {job_title}
Company: {company_name}
Description:
{job_description}

INSTRUCTIONS:
{instructions}

{output_format}"""

    response = client.messages.create(
        model=model,
        max_tokens=max_tokens,
        temperature=temperature,
        messages=[{"role": "user", "content": prompt}],
    )
    return response.content[0].text


def generate_cover_letter(
    client: Anthropic,
    config: dict,
    base_resume: str,
    job_title: str,
    company_name: str,
    job_description: str,
) -> str:
    """Use Claude to generate a tailored cover letter for a specific job."""
    global_rules = config.get("global_instructions", "")
    instructions = config.get("cover_letter_instructions", "Write a professional cover letter.")
    output_format = config.get("cover_letter_output_format", "Return only the cover letter text.")
    model = config.get("model", "claude-sonnet-4-20250514")
    max_tokens = config.get("cover_letter_max_tokens", 1500)
    temperature = config.get("temperature", 1.0)

    prompt = f"""You are a professional cover letter writer. Write a tailored cover letter for the following job application.

GLOBAL RULES (always follow these):
{global_rules}

CANDIDATE'S RESUME:
{base_resume}

TARGET JOB:
Title: {job_title}
Company: {company_name}
Description:
{job_description}

INSTRUCTIONS:
{instructions}

{output_format}"""

    response = client.messages.create(
        model=model,
        max_tokens=max_tokens,
        temperature=temperature,
        messages=[{"role": "user", "content": prompt}],
    )
    return response.content[0].text


# ──────────────────────────────────────────────────────────────────────
# DOCX GENERATION
# ──────────────────────────────────────────────────────────────────────

def sanitize_filename(text: str) -> str:
    """Convert text to a safe filename."""
    text = re.sub(r"[^\w\s-]", "", text)
    text = re.sub(r"\s+", "_", text.strip())
    return text[:50]


def add_inline_runs(paragraph, text: str, base_size: Pt = Pt(10.5)) -> None:
    """Parse inline Markdown (**bold** and *italic*) and add formatted runs to a paragraph."""
    # Split on **bold** and *italic* markers
    parts = re.split(r"(\*\*.*?\*\*|\*.*?\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.size = base_size
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
            run.font.size = base_size
        else:
            run = paragraph.add_run(part)
            run.font.size = base_size


def create_resume_docx(resume_text: str, config: dict, job_title: str, company_name: str, output_path: str) -> None:
    """Create a formatted DOCX resume from Markdown-formatted text."""
    personal = config.get("personal", {})
    name = personal.get("name", "Your Name")
    title = personal.get("title", "Software Engineer")
    location = personal.get("location", "")
    email = personal.get("email", "")
    linkedin = personal.get("linkedin", "")
    github = personal.get("github", "")
    website = personal.get("website", "")

    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10.5)

    # Set narrow margins
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    # Header — Name
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(name)
    name_run.bold = True
    name_run.font.size = Pt(18)

    # Header — Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(title)
    title_run.font.size = Pt(12)

    # Header — Contact
    contact_parts = []
    if location:
        contact_parts.append(location)
    if email:
        contact_parts.append(email)
    contact_line_1 = " | ".join(contact_parts)

    link_parts = []
    if linkedin:
        link_parts.append(linkedin)
    if github:
        link_parts.append(github)
    if website:
        link_parts.append(website)
    contact_line_2 = " | ".join(link_parts)

    contact_text = contact_line_1
    if contact_line_2:
        contact_text += "\n" + contact_line_2

    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_run = contact_para.add_run(contact_text)
    contact_run.font.size = Pt(9)

    # Parse Markdown and add resume sections
    lines = resume_text.strip().split("\n")

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        # ## Section headers
        if stripped.startswith("## "):
            heading_text = stripped[3:].strip()
            heading = doc.add_heading(heading_text, level=2)
            heading.runs[0].font.size = Pt(12)
            continue

        # Bullet points (- item)
        if stripped.startswith("- "):
            bullet_text = stripped[2:].strip()
            para = doc.add_paragraph(style="List Bullet")
            add_inline_runs(para, bullet_text, Pt(10.5))
            para.paragraph_format.space_after = Pt(2)
            para.paragraph_format.space_before = Pt(0)
        # Italic lines (*text*) — used for date ranges
        elif stripped.startswith("*") and stripped.endswith("*") and not stripped.startswith("**"):
            para = doc.add_paragraph()
            run = para.add_run(stripped[1:-1])
            run.italic = True
            run.font.size = Pt(9.5)
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(4)
        # Bold lines (**text**) — used for job titles and skill labels with colons
        elif stripped.startswith("**") and ":" in stripped:
            # Skill line like **Frontend:** React, Next.js, ...
            para = doc.add_paragraph()
            add_inline_runs(para, stripped, Pt(10))
            para.paragraph_format.space_after = Pt(2)
        elif stripped.startswith("**") and stripped.endswith("**"):
            # Job title line like **Senior Dev — Company (Remote)**
            para = doc.add_paragraph()
            run = para.add_run(stripped[2:-2])
            run.bold = True
            run.font.size = Pt(10.5)
            para.paragraph_format.space_after = Pt(0)
        else:
            para = doc.add_paragraph()
            add_inline_runs(para, stripped)
            para.paragraph_format.space_after = Pt(2)

    doc.save(output_path)


def create_cover_letter_docx(cover_letter_text: str, config: dict, job_title: str, company_name: str, output_path: str) -> None:
    """Create a formatted DOCX cover letter."""
    personal = config.get("personal", {})
    name = personal.get("name", "Your Name")

    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Date
    date_para = doc.add_paragraph()
    date_para.add_run(datetime.now().strftime("%B %d, %Y"))
    date_para.paragraph_format.space_after = Pt(12)

    # Recipient
    doc.add_paragraph(f"Re: {job_title} at {company_name}")
    doc.add_paragraph("")

    # Body paragraphs
    paragraphs = cover_letter_text.strip().split("\n\n")
    for para_text in paragraphs:
        clean_text = para_text.strip().replace("\n", " ")
        if clean_text:
            para = doc.add_paragraph(clean_text)
            para.paragraph_format.space_after = Pt(8)

    # Signature
    doc.add_paragraph("")
    sig = doc.add_paragraph()
    sig.add_run("Sincerely,\n").font.size = Pt(11)
    name_run = sig.add_run(name)
    name_run.bold = True
    name_run.font.size = Pt(11)

    doc.save(output_path)


# ──────────────────────────────────────────────────────────────────────
# MAIN
# ──────────────────────────────────────────────────────────────────────

def main() -> None:
    parser = argparse.ArgumentParser(description="Generate tailored resumes and cover letters for top job matches.")
    parser.add_argument("--matches", "-m", default=None, help="Path to best_matches CSV (default: auto-detect latest)")
    parser.add_argument("--jobs", "-j", default=None, help="Path to jobs CSV with full descriptions (default: auto-detect latest)")
    parser.add_argument("--top", "-t", type=int, default=None, help="Number of top jobs to generate for (default: from config)")
    parser.add_argument("--job", "-n", type=int, default=None, help="Generate for a single specific job by rank (e.g. --job 3 for the 3rd match)")
    parser.add_argument("--output-dir", "-d", default=None, help="Output directory (default: output_<timestamp>)")
    parser.add_argument("--config", "-c", default=DEFAULT_CONFIG_PATH, help=f"Path to AI config YAML (default: {DEFAULT_CONFIG_PATH})")
    args = parser.parse_args()

    # Load config
    try:
        config = load_config(args.config)
    except FileNotFoundError:
        return

    # Load base resume
    try:
        base_resume = load_base_resume()
    except FileNotFoundError:
        return

    # Resolve top N (CLI arg overrides config)
    top_n = args.top or config.get("default_top_n", 5)

    # --job takes precedence over --top
    single_job_index = args.job

    # Resolve input files
    matches_file = args.matches or find_latest_matches_csv()
    jobs_file = args.jobs or find_latest_jobs_csv()

    if not matches_file:
        print("❌ No best_matches CSV found. Run the matcher first:")
        print("   python matcher.py")
        return

    if not jobs_file:
        print("❌ No jobs CSV found. Run the scraper first:")
        print("   python scraper.py")
        return

    # Check for API key
    api_key = os.environ.get("ANTHROPIC_API_KEY")
    if not api_key:
        # Try loading from .env file
        env_path = os.path.join(os.path.dirname(__file__), ".env")
        if os.path.exists(env_path):
            with open(env_path) as f:
                for line in f:
                    if line.strip().startswith("ANTHROPIC_API_KEY="):
                        api_key = line.strip().split("=", 1)[1].strip().strip('"').strip("'")
                        break

    if not api_key:
        print("❌ ANTHROPIC_API_KEY not found.")
        print("   Set it as an environment variable:")
        print("     export ANTHROPIC_API_KEY=sk-ant-...")
        print("   Or create a .env file with:")
        print("     ANTHROPIC_API_KEY=sk-ant-...")
        return

    # Output dir name is determined now but created later (after confirmation if --job is used)
    output_dir = args.output_dir or f"output_{datetime.now().strftime('%Y-%m-%d_%H-%M-%S')}"

    model = config.get("model", "claude-sonnet-4-20250514")

    print(f"📄 Resume & Cover Letter Generator")
    print(f"   Config:       {args.config}")
    print(f"   Matches file: {matches_file}")
    print(f"   Jobs file:    {jobs_file}")
    print(f"   Top N:        {top_n}")
    print(f"   Output dir:   {output_dir}")
    print(f"   Model:        {model}")
    print(f"   Temperature:  {config.get('temperature', 1.0)}")

    # Load data
    matches = load_csv(matches_file)
    description_map = build_job_description_map(jobs_file)

    # Get target matches
    if single_job_index is not None:
        if single_job_index < 1 or single_job_index > len(matches):
            print(f"❌ --job {single_job_index} is out of range (matches has {len(matches)} jobs)")
            return
        job = matches[single_job_index - 1]
        job_title = job.get("job_title", "Unknown")
        company_name = job.get("company_name", "Unknown")
        job_link = job.get("job_link", "")
        score = job.get("score", "?")
        print(f"\n   Job #{single_job_index}:")
        print(f"   Title:   {job_title}")
        print(f"   Company: {company_name}")
        print(f"   Score:   {score}")
        print(f"   Link:    {job_link}")

        # Warn if already applied
        if os.path.exists("applications.csv"):
            with open("applications.csv", mode="r", encoding="utf-8") as _f:
                import csv as _csv
                _applied_path = job_link.strip().split("?")[0].rstrip("/")
                for _row in _csv.DictReader(_f):
                    if _row.get("job_link", "").strip().split("?")[0].rstrip("/") == _applied_path:
                        print(f"\n   ⚠️  Already applied on {_row.get('date_applied', '?')} (status: {_row.get('status', '?')})")
                        break

        confirm = input("\n   Continue with this job? [Y/n]: ").strip().lower()
        if confirm not in {"", "y", "yes"}:
            print("   Cancelled.")
            return
        top_matches = [job]
        print(f"\n   Generating documents for job #{single_job_index}...\n")
    else:
        top_matches = matches[:top_n]
        print(f"\n   Generating documents for {len(top_matches)} jobs...\n")

    os.makedirs(output_dir, exist_ok=True)

    client = Anthropic(api_key=api_key)

    for i, match in enumerate(top_matches, start=1):
        job_title = match.get("job_title", "Unknown")
        company_name = match.get("company_name", "Unknown")
        job_link = match.get("job_link", "")
        score = match.get("score", "?")

        # Get full job description from the jobs CSV
        job_description = description_map.get(job_link.strip(), "")
        if not job_description:
            print(f"  ⚠️  [{i}/{len(top_matches)}] No description found for {job_title} @ {company_name}, skipping")
            continue

        safe_company = sanitize_filename(company_name)
        safe_title = sanitize_filename(job_title)
        file_prefix = f"{i:02d}_{safe_company}_{safe_title}"

        print(f"  🔄 [{i}/{len(top_matches)}] {job_title} @ {company_name} (score: {score})")

        # Save job description for interview prep
        jd_path = os.path.join(output_dir, f"{file_prefix}_job_description.md")
        with open(jd_path, "w", encoding="utf-8") as f:
            f.write(f"# {job_title} at {company_name}\n\n")
            f.write(f"**Link:** {job_link}\n\n")
            f.write(f"**Match Score:** {score}\n\n")
            f.write("## Job Description\n\n")
            f.write(job_description)

        # Generate tailored resume
        print(f"     📝 Generating resume...")
        try:
            resume_text = generate_tailored_resume(client, config, base_resume, job_title, company_name, job_description)
            resume_path = os.path.join(output_dir, f"{file_prefix}_resume.docx")
            create_resume_docx(resume_text, config, job_title, company_name, resume_path)
            with open(os.path.join(output_dir, f"{file_prefix}_resume.md"), "w", encoding="utf-8") as f:
                f.write(resume_text)
            print(f"     ✅ Resume saved: {resume_path}")
        except Exception as e:
            print(f"     ❌ Resume generation failed: {e}")

        # Generate cover letter
        print(f"     ✉️  Generating cover letter...")
        try:
            cover_letter_text = generate_cover_letter(client, config, base_resume, job_title, company_name, job_description)
            cl_path = os.path.join(output_dir, f"{file_prefix}_cover_letter.docx")
            create_cover_letter_docx(cover_letter_text, config, job_title, company_name, cl_path)
            with open(os.path.join(output_dir, f"{file_prefix}_cover_letter.md"), "w", encoding="utf-8") as f:
                f.write(cover_letter_text)
            print(f"     ✅ Cover letter saved: {cl_path}")
        except Exception as e:
            print(f"     ❌ Cover letter generation failed: {e}")

        print()

    print(f"{'='*60}")
    print(f"  🎉 Done! Documents saved to {output_dir}/")
    print(f"{'='*60}")


if __name__ == "__main__":
    main()
